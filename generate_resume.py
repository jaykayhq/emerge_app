from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume():
    doc = Document()

    # Define some helper functions for styling
    def add_section_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(46, 116, 181) # Professional blue
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Add a bottom border effectively by underline or just bolding
        # docx doesn't have a simple horizontal line, so we'll just use headers.

    def add_job_title(title, dates):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        p.add_run(f"\t{dates}")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_bullet(text):
        p = doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.left_indent = Pt(20)

    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run("JOSHUA UKPAI")
    name_run.bold = True
    name_run.font.size = Pt(24)

    headline_p = doc.add_paragraph()
    headline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline_run = headline_p.add_run("Web & Android Developer | Project Manager | Market Researcher")
    headline_run.font.size = Pt(12)
    headline_run.italic = True

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.add_run("Port-Harcourt, Rivers State, Nigeria | ")
    contact_p.add_run("joeukpai55@gmail.com | +2349013213794\n")
    contact_p.add_run("LinkedIn Profile: https://www.linkedin.com/in/joshua-ukpai-7194271b8/")

    # Summary
    add_section_header("Professional Summary")
    summary = doc.add_paragraph("Highly motivated Web & Android Developer and Project Manager with expertise in delivering innovative digital solutions. Specialized in market research, software development, and the implementation of AI-driven applications. Proven track record in managing technical projects and developing user-centric habit engines and data analysis tools.")

    # Skills
    add_section_header("Core Competencies")
    skills = [
        "Development: Web & Android Development, AI Integration, Machine Learning.",
        "Management: Project Management, Contract Management, Agile Methodologies.",
        "Tools & Tech: IT Infrastructure, Microsoft Excel, Microsoft Office Suite.",
        "Research: Market Research, Data Analysis, Research Methodologies."
    ]
    for skill in skills:
        add_bullet(skill)

    # Experience
    add_section_header("Professional Experience")

    add_job_title("Computer Technical Specialist Intern", "Jun 2024 – Sep 2024")
    doc.add_paragraph("Petroleum Training Institute (P.T.I.), Nigeria")
    add_bullet("Provided technical support and maintained IT systems.")
    add_bullet("Assisted in software implementation and hardware troubleshooting.")

    add_job_title("Information Technology Specialist Intern", "Mar 2024 – Jun 2024")
    doc.add_paragraph("Admiralty University Of Nigeria (ADUN)")
    add_bullet("Supported IT infrastructure and university-wide digital services.")
    add_bullet("Collaborated on network maintenance and system optimization.")

    add_job_title("Assistant Quantity Surveyor (Intern)", "Jul 2023 – Sep 2023")
    doc.add_paragraph("Monier Construction Company")
    add_bullet("Assisted in contract management and project cost estimation.")
    add_bullet("Utilized Microsoft Excel for detailed project tracking and reporting.")

    # Projects
    add_section_header("Key Projects")

    add_job_title("Emerge — Identity-First Habit Engine", "Dec 2025 – Present")
    add_bullet("Developed an RPG-style habit tracker based on Atomic Habits science.")
    add_bullet("Implemented a 3D-animated avatar system that levels up with real-world user progress.")
    add_bullet("Integrated AI-driven insights to personalize habit loops and identity reinforcement.")

    add_job_title("SocialSense", "Jul 2025 – Present")
    add_bullet("Built a specialized social media analysis application for Nigerian SMEs.")
    add_bullet("Leveraged data analytics to provide actionable growth insights for small businesses.")

    # Education
    add_section_header("Education")
    edu_p = doc.add_paragraph()
    run = edu_p.add_run("Afe Babalola University")
    run.bold = True
    edu_p.add_run("\t2021 – 2025 (Expected)")
    doc.add_paragraph("Bachelor of Engineering - Computer Engineering")
    add_bullet("Leadership: Course Representative, IEEE member.")
    add_bullet("Involvement: Actively participated in IEEE-related events and technical workshops.")

    # Save the document
    doc.save("Joshua_Ukpai_Resume.docx")
    print("Resume generated successfully as Joshua_Ukpai_Resume.docx")

if __name__ == "__main__":
    create_resume()
